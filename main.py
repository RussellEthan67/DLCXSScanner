# Test di Kali
import requests
import msvcrt
import random
import time
import re
import sys
import validators
from pyfiglet import figlet_format
from colorama import Fore, init
from urllib.parse import urlparse
from wafw00f.main import WAFW00F
from scraper import *
from ConnectDatabase import *
from concurrent.futures import ThreadPoolExecutor, as_completed
from fake_useragent import UserAgent
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import smtplib
from email.message import EmailMessage
import mimetypes

init(autoreset=True)

dangerous_characters = [">", "'", '"', "<", "/", ";"]

default_settings = {
    "filename": None,
    "url": None,
    "output": None,
    "threads": None,
    "headers": None,
    "waf": True,
    "custom_waf": None,
    "crawl": False,
    "silent": False
}


def waf_detect(url: str) -> str:
    wafw00f = WAFW00F(url)
    result = wafw00f.identwaf()

    if result:
        result = result[0].lower()
    else:
        return None

    wafs = fetch_names("waf_list.txt")

    for waf in wafs:
        if waf in result:
            return waf

    return None


@staticmethod
def fetch_names(filename: str):
    with open(filename, "r") as waf_list:
        return waf_list.read().split()


def xss(
    filename: str = None,
    url: str = None,
    output: str = None,
    threads: int = None,
    headers: list = None,
    waf: bool = False,
    custom_waf: str = None,
    crawl: bool = False,
    silent: bool = False
):
    if headers is not None:
        try:
            headers = headerParser(headers)
        except AttributeError:
            headers = headerParser(headers)

    if crawl is True:
        filename = "link.txt"

    urls = []
    results = []

    if url is not None and filename is None:
        res = scanner(url, headers, waf, custom_waf, silent)
        results.append(res)
        print(Fore.YELLOW + "\nScanning process complete. Writing results now....")
        if results:
            write(output, results)
        return None
    elif crawl is True:
        output = crawling(url)
        print(Fore.GREEN + "Crawling has already completed.")
        urls = read(output)
    else:
        urls = read(filename)

    if urls:
        print(Fore.YELLOW +
              f"There are {len(urls)} URL in this list. Working on it....")
        if threads is not None and threads > 0:
            with ThreadPoolExecutor(max_workers=threads) as executor:
                future_to_url = {executor.submit(
                    scanner, url, headers, waf, custom_waf, silent): url for url in urls}
                for future in as_completed(future_to_url):
                    url = future_to_url[future]
                    try:
                        res = future.result()
                        print(Fore.YELLOW + f"Working on {url}...")
                    except Exception as exc:
                        print(Fore.RED + '%r generated an exception: %s' %
                              (url, exc))
                    else:
                        results.append(res)
        elif threads == 0:
            print(Fore.RED + "Thread cannot be set to 0")
        else:
            for url in urls:
                print(Fore.YELLOW + f"Working on {url}...")
                res = scanner(url, headers, waf, custom_waf, silent)
                results.append(res)
    print(Fore.YELLOW + "Scanning process complete. Writing results now....")
    write(output, results)


def headerParser(input_list: list) -> dict:
    outputDictionary = {}

    for item in input_list:
        key_value = re.split(r":\s*", item, 1)

        if len(key_value) == 2:
            key, value = key_value
            outputDictionary[key.strip()] = value.strip()

    return outputDictionary


def read(filename: str) -> list:
    with open(filename, "r") as file:
        lines = (line.rstrip() for line in file if line.strip())
        return list(lines)


def send_email_with_attachment(sender_email, receiver_email, subject, body, file_path, smtp_server, smtp_port, login, password):
    # Create the email message
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.set_content(body)

    # Determine the MIME type of the file
    mime_type, _ = mimetypes.guess_type(file_path)
    mime_type, mime_subtype = mime_type.split('/')

    # Attach the file
    with open(file_path, 'rb') as file:
        msg.add_attachment(file.read(),
                           maintype=mime_type,
                           subtype=mime_subtype,
                           filename=file_path)

    # Connect to the SMTP server and send the email
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()  # Secure the connection
            server.login(login, password)
            server.send_message(msg)
            print(Fore.GREEN + f"Email sent to {receiver_email}")
    except Exception as e:
        print(Fore.RED + f"Failed to send email: {e}")


def write_to_docx(output: str, results: list):
    """
    Create a .docx report from XSS scanner results.

    Parameters:
    scan_results (list): List of dictionaries with scan results.
    output (str): Path to save the .docx file.
    """
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading(level=1)
    title_run = title.add_run('XSS Scanner Report')
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add summary
    doc.add_heading('Summary', level=2)
    total_scans = len(results)
    unsafe_scans = len(
        [result for result in results if not result["is_safe"]])
    summary = f"Total URLs scanned: {total_scans}\n" \
              f"Unsafe URLs found: {unsafe_scans}\n" \
              f"Safe URLs found: {total_scans - unsafe_scans}"
    doc.add_paragraph(summary)

    # Add detailed results
    for result in results:
        url_heading = doc.add_heading(f'Target URL: {result["url"]}', level=2)
        url_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if result["is_safe"]:
            safe_paragraph = doc.add_paragraph('This URL is safe.')
            safe_paragraph.italic = True
        else:
            not_safe_paragraph = doc.add_paragraph('This URL is not safe!')
            not_safe_paragraph.bold = True
            for vuln in result["vulnerabilities"]:
                vuln_heading = doc.add_heading(level=3)
                vuln_heading_run = vuln_heading.add_run(
                    f"Vulnerability: {vuln['type']}")
                vuln_heading_run.bold = True
                vuln_heading_run.font.color.rgb = RGBColor(
                    255, 0, 0)  # Red color for vulnerability type

                # Create a paragraph and add bold text for "Payload Used:"
                payload_paragraph = doc.add_paragraph()
                payload_run = payload_paragraph.add_run("Payload Used: ")
                payload_run.bold = True
                payload_paragraph.add_run(f"{vuln['payload']}")

                full_url_paragraph = doc.add_paragraph()
                full_url_run = full_url_paragraph.add_run(
                    f"Full URL: {vuln['full_url']}")
                full_url_run.font.color.rgb = RGBColor(
                    0, 0, 255)  # Blue color for full URL
                screenshot_path = take_screenshot(vuln['full_url'])
                if screenshot_path:
                    doc.add_paragraph("Screenshot:")
                    doc.add_picture(screenshot_path, width=Inches(6.0))
                else:
                    doc.add_paragraph(
                        "Screenshot could not be taken.", style='Normal').italic = True

    # Save the document
    doc.save(output)
    print(Fore.GREEN + f"Report saved as {output}")


def is_valid_email(email):
    # Regular expression for validating an Email
    regex = r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$'
    return re.match(regex, email) is not None


def write(output: str, results: list) -> None:
    """
    Writes the result to a specified output file or prints to the console if no file is specified.

    Parameters:
    output (str): The file path to write to. Supports .txt and .docx extensions.
    result (list): The list of results to write or print.
    """
    if len(results) == 0:
        print(Fore.GREEN + "URL seems to be safe." + Fore.RESET)
        return None

    if not output:
        for result in results:
            for vuln in result["vulnerabilities"]:
                print("This link seems to work: " + Fore.RED +
                      vuln['full_url'] + Fore.RESET)
        return None

    if output.endswith(".txt"):
        try:
            with open(output, "a") as output_file:
                for result in results:
                    output_file.write(f"{result['full_url']}\n")
            print(Fore.GREEN + "Writing completed successfully.")
        except IOError as e:
            print(Fore.RED +
                  f"An error occurred while writing to the file: {e}")

    elif output.endswith(".docx"):
        print(Fore.YELLOW + f"Writing {output} file...")
        write_to_docx(output, results)
        print(Fore.GREEN + "Writing completed successfully.")
        inp = input("Do you want the report sent to your email?(Yes/No): ")
        if inp.lower() == "no":
            print("You can find the .docx file in this application directory")
        elif inp.lower() == "yes":
            while True:
                receiver_email = input("Insert your email address here: ")
                if not is_valid_email(receiver_email):
                    print(Fore.RED + "Invalid email address!")
                else:
                    break
            sender_email = 'dlcxss@outlook.com'
            subject = 'XSS Scanner Report'
            body = 'Please find the attached XSS scanner report.'
            smtp_server = 'smtp.office365.com'
            smtp_port = 587  # For TLS
            login = 'dlcxss@outlook.com'
            password = 'canner789'
            send_email_with_attachment(
                sender_email, receiver_email, subject, body, output, smtp_server, smtp_port, login, password)

    else:
        print(Fore.RED + f"Unsupported file format: {output}")


def replace(url: str, param_name: str, value: str) -> str:
    return re.sub(f"{param_name}=([^&]+)", f"{param_name}={value}", url)


def bubble_sort(arr: list) -> list:
    a = 0
    keys = []

    for i in arr:
        for j in i:
            keys.append(j)

    while a < len(keys) - 1:
        b = 0
        while b < len(keys) - 1:
            d1 = arr[b]
            d2 = arr[b + 1]
            if len(d1[keys[b]]) < len(d2[keys[b + 1]]):
                d = d1
                arr[b] = arr[b + 1]
                arr[b + 1] = d
                z = keys[b + 1]
                keys[b + 1] = keys[b]
                keys[b] = z
            b += 1
        a += 1

    return arr


def crawling(url: str) -> None:
    driver, info = driver_init(url)
    crawl_link(url, driver, info)
    driver.quit()
    filename = "link.txt"
    with open(filename, 'a+') as f:
        f.seek(0)
        lines = f.readlines()
        for links in deep_links:
            if links + '\n' not in lines:
                f.write(links + '\n')
    return filename


def parameters(url: str) -> list:
    param_names = []
    params = urlparse(url).query
    params = params.split("&")

    if len(params) == 1:
        params = params[0].split("=")
        param_names.append(params[0])
    else:
        for param in params:
            param = param.split("=")
            param_names.append(param[0])

    return param_names


def parserer(url: str, param_name: str, value: str) -> dict:
    final_parameters = {}
    parsed_data = urlparse(url)
    params = parsed_data.query
    params = params.split("&")

    if len(params) == 1:
        params = params[0].split("=")
        if len(params) >= 2:
            final_parameters[params[0]] = params[1]
    else:
        for param in params:
            param = param.split("=")
            final_parameters[param[0]] = param[1]

    final_parameters[param_name] = value

    return final_parameters


def validator(arr: list, param_name: str, url: str, headers: dict = None) -> dict:
    dic = {param_name: []}

    for data in arr:
        final_parameters = parserer(url, param_name, data + "randomstring")
        new_url = (
            urlparse(url).scheme
            + "://"
            + urlparse(url).hostname
            + "/"
            + urlparse(url).path
        )

        if headers:
            response = requests.get(
                new_url,
                params=final_parameters,
                headers=headers,
                verify=False,
            ).text
        else:
            response = requests.get(
                new_url, params=final_parameters, verify=False).text

        if data + "randomstring" in response:
            dic[param_name].append(data)

    return dic


def fuzzer(url: str, headers: dict = None) -> list:
    data = []
    params = parameters(url)

    if "" in params and len(params) == 1:
        return None

    for param in params:
        out = validator(dangerous_characters, param, url, headers)
        data.append(out)

    return bubble_sort(data)


def filter_payload(arr: list, firewall: str = None) -> list:
    payload_list = []
    dbs = open("payloads.json")
    dbs = json.load(dbs)
    new_dbs = []

    if firewall:
        for i in range(0, len(dbs)):
            if dbs[i]["waf"] == firewall:
                new_dbs.append(dbs[i])

        if not new_dbs:
            return None
    else:
        for i in range(0, len(dbs)):
            if not dbs[i]["waf"]:
                new_dbs.append(dbs[i])

    dbs = new_dbs

    for char in arr:
        for payload in dbs:
            attributes = payload["Attribute"]
            if char in attributes:
                payload["count"] += 1

    def fun(e):
        return e["count"]

    dbs.sort(key=fun, reverse=True)

    for payload in dbs:
        payload_list.append(payload["Payload"])

    return payload_list


def scanner(url: str, headers: dict = None, waf: bool = False, custom_waf: str = None, silent: bool = False) -> list:
    scan_result = {}

    if waf is True:
        firewall = waf_detect(url)
        if not firewall:
            firewall = None
    elif custom_waf is not None:
        firewall = custom_waf
    else:
        firewall = None

    out = fuzzer(url, headers)
    for data in out:
        for key in data:
            payload_list = filter_payload(data[key], firewall)
        for payload in payload_list:
            data = parserer(url, key, payload)
            parsed_data = urlparse(url)
            new_url = parsed_data.scheme + "://" + parsed_data.netloc + parsed_data.path
            if headers:
                response = requests.get(
                    new_url, params=data, headers=headers, verify=False
                ).text
            elif silent is True:
                ua = UserAgent(browsers=['firefox', 'chrome'])
                streamlined_uas = ua.random
                request_headers = {
                    'user-agent': streamlined_uas
                }
                response = requests.get(
                    new_url, params=data, headers=request_headers, verify=False
                ).text
                time.sleep(random.uniform(0.5, 5))
            else:
                response = requests.get(
                    new_url, params=data, verify=False).text
            if payload in response:
                full_url = replace(url, key, payload)
                scan_result = {
                    "url": url,
                    "is_safe": False,
                    "vulnerabilities": [
                        {"type": "XSS", "payload": payload, "full_url": full_url}
                    ]
                }
                return scan_result

    return None


def default(filename: str, url: str, user_input: str):
    if filename is not None:
        default_settings['filename'] = filename

    if url is not None:
        default_settings['url'] = url

    user_input = input(
        "Do you want to use default settings or custom settings? : ")
    if user_input.lower() in ['default', 'custom']:
        if user_input == 'default':
            xss(default_settings['filename'], default_settings['url'],
                default_settings['output'], default_settings['threads'], default_settings['headers'], default_settings['waf'], default_settings['custom_waf'], default_settings['crawl'])
        elif user_input == 'custom':
            while True:
                print("\n")
                for key, value in default_settings.items():
                    if user_input == "1":
                        if key == "crawl":
                            pass
                    else:
                        print(f"{key.capitalize()} : {str(value)}")
                print("\n")
                key_input = input(
                    "Type settings you want to change (Type save to save settings, Type exit to go back to main menu): ")
                if key_input in default_settings:
                    if user_input == "1":
                        if key_input == "crawl":
                            print(
                                Fore.RED + "We currently do not support multiple links crawling!")
                    value_input = input("Set the value : ")
                    if isinstance(default_settings[key_input], bool):
                        if value_input.lower() == 'true':
                            default_settings[key_input] = True
                        elif value_input.lower() == 'false':
                            default_settings[key_input] = False
                        else:
                            print(Fore.RED +
                                  "Invalid value for boolean setting. Please enter True or False.")
                    else:
                        default_settings[key_input] = value_input
                elif key_input.lower() == "save":
                    print(
                        "Do you want to continue scanning with these settings? (Press Y to continue)")
                    key = msvcrt.getch().decode('utf-8').upper()
                    if key == 'Y':
                        break
                elif key_input.lower() == "exit":
                    return
                else:
                    print("Invalid settings")
            xss(default_settings['filename'], default_settings['url'],
                default_settings['output'], default_settings['threads'], default_settings['headers'], default_settings['waf'], default_settings['custom_waf'], default_settings['crawl'], default_settings['silent'])


def main_ui():
    filename = None
    url = None

    while True:
        print("\n1. File Scan")
        print("2. Single URL Scan")
        print("3. Contribute/Download WAF Payload")
        print("4. Exit\n")
        user_input = input("Select Option: ")

        if user_input == "1":
            filename = input("Input your file name to scan: ")
            default(filename, url, user_input)
        elif user_input == "2":
            while True:
                url = input("Input your URL link to scan: ")
                if not validators.url(url):
                    print(Fore.RED + "Invalid URL! Please enter again.")
                else:
                    break
            default(filename, url, user_input)
        elif user_input == "3":
            start()
        elif user_input == "4":
            print("Exiting the program...\n")
            break
        else:
            print(Fore.RED + "Invalid Option!")


# Creating UI
print(Fore.BLUE + figlet_format("DLC XSS", font="slant") + Fore.RESET)
print(
    Fore.LIGHTBLUE_EX +
    figlet_format("Deep Link Community XSS Scanner", font="digital") +
    Fore.RESET
)
# Create the parser
parser = argparse.ArgumentParser()

# Add the arguments
parser.add_argument(
    "-f", "--filename", help="specify Filename to scan. Eg: urls.txt etc"
)
parser.add_argument(
    "-u", "--url", help="scan a single URL. Eg: http://example.com/?id=2"
)
parser.add_argument(
    "-o", "--output", help="filename to store output. Eg: result.txt")
parser.add_argument(
    "-t", "--threads", help="no of threads to send concurrent requests(Max: 10)"
)
parser.add_argument("-H", "--headers", help="specify Custom Headers")
parser.add_argument(
    "--waf",
    action="store_true",
    help="detect web application firewall and then test payloads",
)
parser.add_argument("-w", "--custom_waf",
                    help="use specific payloads related to W.A.F")
parser.add_argument("--crawl", action="store_true", help="crawl then find xss")


# Parse the arguments
args = parser.parse_args()

# Check if no arguments were passed
if len(sys.argv) != 1:
    # Store the arguments
    filename = args.filename
    url = args.url
    output = args.output
    threads = args.threads
    headers = args.headers
    waf = args.waf
    custom_waf = args.custom_waf
    crawl = args.crawl
    xss(filename, url, output, threads, headers, waf, custom_waf, crawl)
else:
    main_ui()
